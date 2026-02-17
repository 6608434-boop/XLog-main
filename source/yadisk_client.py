import os
import yadisk
import io
from datetime import datetime
from typing import Optional, List
import tempfile
from docx import Document

from .logger import logger

# Пытаемся импортировать определение кодировки (оставим для TXT файлов)
try:
    from charset_normalizer import from_bytes

    CHARSET_DETECT_AVAILABLE = True
    logger.info("charset-normalizer loaded successfully")
except ImportError:
    CHARSET_DETECT_AVAILABLE = False
    logger.warning("charset-normalizer not installed, using fallback encoding detection")


class YandexDiskClient:
    def __init__(self, token: str, root_folder: str = "XLog"):
        """
        Инициализация клиента Яндекс.Диска.

        Args:
            token: OAuth-токен для доступа к Диску
            root_folder: Корневая папка проекта на Диске
        """
        self.client = yadisk.Client(token=token)
        self.root_folder = root_folder

        # Проверяем, что токен рабочий
        if not self.client.check_token():
            logger.error("Invalid Yandex Disk token")
            raise ValueError("Invalid Yandex Disk token")

        logger.info("Connected to Yandex Disk")

    def ensure_folder_exists(self, remote_path: str) -> bool:
        """
        Убеждается, что папка существует. Если нет — создаёт все промежуточные папки.

        Args:
            remote_path: Путь к папке на Яндекс.Диске (относительно корня)

        Returns:
            True если папка существует или была создана
        """
        full_path = f"/{self.root_folder}/{remote_path}"
        try:
            # Разбиваем путь на части и создаём каждую папку по очереди
            current_path = f"/{self.root_folder}"
            parts = remote_path.split('/')

            for part in parts:
                if not part:  # Пропускаем пустые части
                    continue
                current_path += f"/{part}"
                if not self.client.exists(current_path):
                    self.client.mkdir(current_path)
                    logger.debug(f"Created directory: {current_path}")

            return True

        except Exception as e:
            logger.error(f"Failed to ensure folder {full_path}: {e}")
            return False

    def get_daily_log_path(self, profile_name: str, date: datetime) -> str:
        """
        Формирует путь к файлу лога за указанную дату.

        Returns:
            Полный путь к файлу лога
        """
        return f"{profile_name}/logs/{date.year}/{date.month:02d}/{date.day:02d}/log.txt"

    def read_docx(self, remote_path: str) -> Optional[str]:
        """
        Читает содержимое DOCX файла с Яндекс.Диска.

        Args:
            remote_path: Путь к файлу на Яндекс.Диске (относительно корня)

        Returns:
            Текст из документа или None
        """
        full_path = f"/{self.root_folder}/{remote_path}"
        temp_file = None

        try:
            if not self.client.exists(full_path):
                return None

            # Скачиваем файл во временный
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tf:
                temp_file = tf.name
                self.client.download(full_path, temp_file)

            # Открываем через python-docx
            doc = Document(temp_file)

            # Собираем весь текст из параграфов
            text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])

            logger.info(f"✅ Successfully read DOCX {remote_path} ({len(text)} chars)")
            return text

        except Exception as e:
            logger.error(f"Failed to read DOCX {full_path}: {e}")
            return None

        finally:
            if temp_file and os.path.exists(temp_file):
                os.unlink(temp_file)

    def write_docx(self, remote_path: str, content: str) -> bool:
        """
        Создаёт или перезаписывает DOCX файл на Яндекс.Диске.

        Args:
            remote_path: Путь к файлу на Яндекс.Диске (относительно корня)
            content: Текст для записи в документ

        Returns:
            True если успешно
        """
        full_path = f"/{self.root_folder}/{remote_path}"
        temp_file = None

        try:
            # Убеждаемся, что папка для файла существует
            folder_path = '/'.join(remote_path.split('/')[:-1])
            if folder_path:
                self.ensure_folder_exists(folder_path)

            # Создаём временный DOCX файл
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tf:
                temp_file = tf.name

                # Создаём новый документ
                doc = Document()

                # Добавляем текст (разбиваем по строкам)
                for line in content.split('\n'):
                    doc.add_paragraph(line)

                # Сохраняем
                doc.save(temp_file)

            # Загружаем на Диск
            self.client.upload(temp_file, full_path, overwrite=True)

            logger.info(f"✅ Successfully wrote DOCX {remote_path} ({len(content)} chars)")
            return True

        except Exception as e:
            logger.error(f"Failed to write DOCX {full_path}: {e}")
            return False

        finally:
            if temp_file and os.path.exists(temp_file):
                os.unlink(temp_file)

    def read_file(self, remote_path: str) -> Optional[str]:
        """
        Универсальное чтение файла. Если файл .docx — использует read_docx.
        Если .txt — использует старое чтение с определением кодировки.
        """
        if remote_path.lower().endswith('.docx'):
            return self.read_docx(remote_path)
        else:
            return self._read_text_file(remote_path)

    def _read_text_file(self, remote_path: str) -> Optional[str]:
        """
        Читает текстовый файл с автоопределением кодировки.
        """
        full_path = f"/{self.root_folder}/{remote_path}"
        temp_file = None

        try:
            if not self.client.exists(full_path):
                return None

            with tempfile.NamedTemporaryFile(mode='rb', delete=False) as tf:
                temp_file = tf.name
                self.client.download(full_path, temp_file)

            with open(temp_file, 'rb') as f:
                raw_data = f.read()

            if CHARSET_DETECT_AVAILABLE and raw_data:
                try:
                    result = from_bytes(raw_data).best()
                    if result:
                        content = str(result)
                        if content and content.startswith('\ufeff'):
                            content = content[1:]
                        return content
                except Exception:
                    pass

            # Fallback
            return self._fallback_decode(raw_data, remote_path)

        except Exception as e:
            logger.error(f"Failed to read {full_path}: {e}")
            return None

        finally:
            if temp_file and os.path.exists(temp_file):
                os.unlink(temp_file)

    def _fallback_decode(self, raw_data: bytes, remote_path: str) -> Optional[str]:
        """Запасной метод с перебором кодировок"""
        encodings = ['utf-8', 'windows-1251', 'koi8-r', 'cp866', 'iso-8859-5']

        for encoding in encodings:
            try:
                content = raw_data.decode(encoding)
                logger.info(f"Fallback: {remote_path} decoded as {encoding}")
                if content and content.startswith('\ufeff'):
                    content = content[1:]
                return content
            except UnicodeDecodeError:
                continue

        try:
            return raw_data.decode('utf-8', errors='ignore')
        except:
            return None

    def write_file(self, remote_path: str, content: str) -> bool:
        """
        Универсальная запись файла. Если путь заканчивается на .docx — создаёт DOCX.
        Иначе — обычный текст.
        """
        if remote_path.lower().endswith('.docx'):
            return self.write_docx(remote_path, content)
        else:
            return self._write_text_file(remote_path, content)

    def _write_text_file(self, remote_path: str, content: str) -> bool:
        """Записывает текстовый файл в UTF-8"""
        full_path = f"/{self.root_folder}/{remote_path}"
        temp_file = None

        try:
            folder_path = '/'.join(remote_path.split('/')[:-1])
            if folder_path:
                self.ensure_folder_exists(folder_path)

            with tempfile.NamedTemporaryFile(mode='w', encoding='utf-8', delete=False) as tf:
                temp_file = tf.name
                tf.write(content)
                tf.flush()

            self.client.upload(temp_file, full_path, overwrite=True)
            logger.debug(f"Written to {full_path}")
            return True

        except Exception as e:
            logger.error(f"Failed to write to {full_path}: {e}")
            return False

        finally:
            if temp_file and os.path.exists(temp_file):
                os.unlink(temp_file)

    def list_files(self, remote_path: str) -> List[str]:
        """
        Возвращает список файлов в папке.

        Args:
            remote_path: Путь к папке на Яндекс.Диске (относительно корня)

        Returns:
            Список имён файлов
        """
        full_path = f"/{self.root_folder}/{remote_path}"
        try:
            if not self.client.exists(full_path):
                return []
            return [item.name for item in self.client.listdir(full_path)]
        except Exception as e:
            logger.error(f"Failed to list {full_path}: {e}")
            return []